#!/usr/bin/env python3
# flake8: noqa: E501
"""Executive introduction to inDoc — original figures, launch illustrations, tight Word layout."""
from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import ImageDraw

ROOT = Path(__file__).resolve().parents[1]
ASSET = ROOT / "private-docs/gtm/assets/introducing-indoc"
OUT = ROOT / "private-docs/gtm/introducing-indoc.docx"
MD = ROOT / "private-docs/gtm/introducing-indoc.md"

_spec = importlib.util.spec_from_file_location(
    "gdoc", ROOT / "tools/build_graphrag_article_docx.py"
)
g = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(g)

S, PRIMARY, PRIMARY_50, ORANGE, ORANGE_50 = (
    g.S,
    g.PRIMARY,
    g.PRIMARY_50,
    g.ORANGE,
    g.ORANGE_50,
)
TEAL, TEAL_50, PURPLE, PURPLE_50 = g.TEAL, g.TEAL_50, g.PURPLE, g.PURPLE_50
GREEN, GREEN_50, NAVY, TEXT, MUTED, LINE, PAPER = (
    g.GREEN,
    g.GREEN_50,
    g.NAVY,
    g.TEXT,
    g.MUTED,
    g.LINE,
    g.PAPER,
)
C_PRIMARY, C_NAVY, C_TEXT, C_MUTED = g.C_PRIMARY, g.C_NAVY, g.C_TEXT, g.C_MUTED


def save(img, name: str) -> Path:
    from PIL import Image

    rgb = Image.new("RGB", img.size, g.CANVAS)
    rgb.paste(img, mask=img.split()[-1])
    w = 2400
    rgb = rgb.resize(
        (w, max(1, int(rgb.height * w / rgb.width))), Image.Resampling.LANCZOS
    )
    out = ASSET / name
    ASSET.mkdir(parents=True, exist_ok=True)
    rgb.save(out, "PNG", optimize=True, dpi=(300, 300))
    return out


def fig_rag_vs_agent() -> Path:
    img = g.canvas(1400, 560)
    d0 = ImageDraw.Draw(img)
    g.fig_label(
        d0,
        "Figure 1  ·  Retrieval Versus Investigation",
        "RAG Retrieves Once. inDoc Investigates Until It Has Evidence.",
    )
    g.plate(img, (S(40), S(100), S(680), S(528)), 16)
    d = ImageDraw.Draw(img)
    g.pill(d, (S(64), S(124), S(200), S(152)), ORANGE_50, "RAG", ORANGE, 11)
    g.txt(d, (S(64), S(172)), "RAG", 22, True, NAVY)
    g.wrapped(
        d, S(64), S(214), "Retrieve, stuff context, answer.", S(580), 16, True, ORANGE
    )
    steps = [
        "One retrieval",
        "One generation",
        "A fixed pipeline",
        "The model does not choose the next step",
    ]
    for i, s in enumerate(steps):
        y = S(280) + i * S(52)
        g.draw_rr(d, (S(64), y, S(656), y + S(42)), 8, (245, 247, 250), LINE)
        g.txt(d, (S(84), y + S(21)), f"{i+1:02d}   {s}", 13.5, True, MUTED, "lm")
    g.plate(img, (S(720), S(100), S(1360), S(528)), 16)
    d = ImageDraw.Draw(img)
    g.pill(d, (S(744), S(124), S(860), S(152)), PRIMARY_50, "Now", PRIMARY, 11)
    g.txt(d, (S(744), S(172)), "Agent", 22, True, NAVY)
    g.wrapped(
        d,
        S(744),
        S(214),
        "Plan, act, observe, re-plan, answer.",
        S(580),
        16,
        True,
        PRIMARY,
    )
    steps = [
        "Chooses each tool call",
        "Chains steps from evidence",
        "Stops when it has enough",
        "Returns a full trace",
    ]
    for i, s in enumerate(steps):
        y = S(280) + i * S(52)
        g.draw_rr(d, (S(744), y, S(1336), y + S(42)), 8, PRIMARY_50, PRIMARY)
        g.txt(d, (S(764), y + S(21)), f"{i+1:02d}   {s}", 13.5, True, NAVY, "lm")
    return save(img, "fig-rag-vs-agent.png")


def fig_q3_trace() -> Path:
    img = g.canvas(1400, 620)
    d0 = ImageDraw.Draw(img)
    g.fig_label(
        d0,
        "Figure 2  ·  Investigation Trace",
        "Goal: What was Q3 revenue and its growth rate?",
    )
    rows = [
        ("01", "Think", "Find the financial report first.", ORANGE, ORANGE_50),
        ("01", "Act", "Search documents  ·  revenue", GREEN, GREEN_50),
        ("01", "Observe", "Q3 Financial Report", PRIMARY, PRIMARY_50),
        ("02", "Think", "Read it for the exact figure.", ORANGE, ORANGE_50),
        ("02", "Act", "Read document  ·  9f6c…", GREEN, GREEN_50),
        ("02", "Observe", "$4.2M, up 18% YoY", PRIMARY, PRIMARY_50),
        ("03", "Finish", "Q3 2025 revenue was $4.2M, up 18% YoY", TEAL, TEAL_50),
    ]
    for i, (n, kind, line, accent, wash) in enumerate(rows):
        y = S(100) + i * S(68)
        g.plate(img, (S(40), y, S(1360), y + S(60)), 10)
        d = ImageDraw.Draw(img)
        g.pill(d, (S(56), y + S(14), S(118), y + S(46)), wash, n, accent, 11)
        g.pill(d, (S(132), y + S(14), S(268), y + S(46)), accent, kind, PAPER, 11)
        g.txt(d, (S(292), y + S(30)), line, 16, True, NAVY, "lm")
    return save(img, "fig-q3-trace.png")


def fig_loop() -> Path:
    img = g.canvas(1400, 520)
    d0 = ImageDraw.Draw(img)
    g.fig_label(
        d0,
        "Figure 3  ·  The Loop",
        "The model chooses each move. Scope checks every tool. The trace returns with the answer.",
    )
    nodes = [
        (48, "You", "A goal, in plain language", PRIMARY),
        (328, "Agent", "Plan the next step", ORANGE),
        (608, "LLM", "Thought plus action", TEAL),
        (888, "Tools", "RBAC / ABAC filtered", PURPLE),
        (1168, "You", "Answer plus trace", GREEN),
    ]
    d = ImageDraw.Draw(img)
    for i, (x, title, sub, accent) in enumerate(nodes):
        g.plate(img, (S(x), S(160), S(x + 220), S(360)), 14)
        dd = ImageDraw.Draw(img)
        g.draw_rr(dd, (S(x), S(160), S(x + 220), S(168)), 0, accent)
        g.txt(dd, (S(x + 16), S(196)), f"{i+1:02d}", 12, True, accent)
        g.txt(dd, (S(x + 16), S(232)), title, 18, True, NAVY)
        g.wrapped(dd, S(x + 16), S(276), sub, S(188), 13, False, MUTED)
        if i < 4:
            dd.polygon(
                [(S(x + 228), S(250)), (S(x + 248), S(260)), (S(x + 228), S(270))],
                fill=PRIMARY + (180,),
            )
    g.txt(
        d,
        (S(48), S(400)),
        "Bounded: six steps by default, twelve hard cap, two deep dives. Then finish.",
        13,
        False,
        MUTED,
    )
    g.txt(
        d,
        (S(48), S(436)),
        "Live stream: start  ·  step  ·  final  ·  error.  POST /api/v1/agent/run and /agent/stream.",
        13,
        False,
        MUTED,
    )
    return save(img, "fig-loop.png")


def fig_tools() -> Path:
    img = g.canvas(1400, 620)
    d0 = ImageDraw.Draw(img)
    g.fig_label(
        d0,
        "Figure 4  ·  Agent Tools",
        "Six tools. Every call is access-checked. Nothing outside effective scope.",
    )
    tools = [
        (40, "List Documents", "See what is in scope before acting.", PRIMARY),
        (500, "Search Documents", "Hybrid keyword and semantic search.", ORANGE),
        (960, "Read Document", "Pull full text to extract facts.", TEAL),
        (40, "Summarize Document", "Get the gist of one long file.", PURPLE),
        (500, "Compare Documents", "Find what changed across files.", GREEN),
        (960, "Finish", "Return the grounded answer.", PRIMARY),
    ]
    for i, (x, title, body, accent) in enumerate(tools):
        y = S(108) if i < 3 else S(340)
        g.plate(img, (S(x), y, S(x + 420), y + S(200)), 14)
        d = ImageDraw.Draw(img)
        g.draw_rr(d, (S(x), y, S(x + 8), y + S(200)), 0, accent)
        g.txt(
            d,
            (S(x + 28), y + S(28)),
            f"{(i%3)+1 if i<3 else i+1:02d}",
            12,
            True,
            accent,
        )
        g.txt(d, (S(x + 28), y + S(64)), title, 16, True, NAVY)
        g.wrapped(d, S(x + 28), y + S(110), body, S(364), 13.5, False, TEXT)
    return save(img, "fig-tools.png")


def fig_pipeline() -> Path:
    img = g.canvas(1400, 480)
    d0 = ImageDraw.Draw(img)
    g.fig_label(
        d0,
        "Figure 5  ·  Ingest",
        "Virus scan, extract, embed, dual index — off the request path.",
    )
    steps = [
        ("01", "Upload", "PDF, Word, Excel, PowerPoint, images, mail", PRIMARY),
        ("02", "Virus Scan", "ClamAV, YARA, integrity", ORANGE),
        ("03", "Extract", "Text and OCR", TEAL),
        ("04", "Embed", "384-d cosine vector", PURPLE),
        ("05", "Index", "Elasticsearch and Qdrant", GREEN),
        ("06", "Ready", "Scoped search and chat", PRIMARY),
    ]
    for i, (n, title, sub, accent) in enumerate(steps):
        x = 40 + i * 226
        g.plate(img, (S(x), S(120), S(x + 210), S(380)), 12)
        d = ImageDraw.Draw(img)
        g.draw_rr(d, (S(x), S(120), S(x + 210), S(128)), 0, accent)
        g.txt(d, (S(x + 16), S(156)), n, 13, True, accent)
        g.txt(d, (S(x + 16), S(200)), title, 16, True, NAVY)
        g.wrapped(d, S(x + 16), S(250), sub, S(178), 12.5, False, MUTED)
        if i < 5:
            d.polygon(
                [(S(x + 214), S(240)), (S(x + 226), S(248)), (S(x + 214), S(256))],
                fill=PRIMARY + (170,),
            )
    d = ImageDraw.Draw(img)
    g.txt(
        d,
        (S(48), S(420)),
        "Celery workers. Postgres record. Local and S3 storage. Progress over the wire.",
        13,
        False,
        MUTED,
    )
    return save(img, "fig-pipeline.png")


def fig_three_pillars() -> Path:
    img = g.canvas(1400, 620)
    d0 = ImageDraw.Draw(img)
    g.fig_label(
        d0, "Figure  ·  The Product", "Agentic AI, hybrid search, and compliance."
    )
    cols = [
        (
            40,
            ORANGE,
            ORANGE_50,
            "01",
            "Agentic AI",
            "Investigation",
            "The agent plans, searches, reads, and finishes — six tools, a live trace, and a step budget.",
        ),
        (
            496,
            PRIMARY,
            PRIMARY_50,
            "02",
            "Hybrid Search",
            "Retrieval",
            "Elasticsearch holds the words. Qdrant holds meaning. Scores fuse at equal weight. Scope is applied before any result leaves. Chat, agent, and tools share the path.",
        ),
        (
            952,
            TEAL,
            TEAL_50,
            "03",
            "Compliance",
            "Control",
            "RBAC and ABAC. HIPAA and PCI modes. Audit and SIEM. Local model first. Restricted text does not appear in an Internal answer.",
        ),
    ]
    for x, accent, wash, num, title, dek, body in cols:
        g.plate(img, (S(x), S(108), S(x + 432), S(588)), 16)
        d = ImageDraw.Draw(img)
        g.draw_rr(d, (S(x), S(108), S(x + 432), S(116)), 0, accent)
        g.pill(d, (S(x + 20), S(136), S(x + 90), S(168)), wash, num, accent, 12)
        g.txt(d, (S(x + 20), S(196)), title, 26, True, NAVY)
        g.txt(d, (S(x + 20), S(248)), dek, 15, True, accent)
        g.wrapped(d, S(x + 20), S(300), body, S(392), 14.5, False, TEXT, 1.35)
    return save(img, "fig-three-pillars.png")


def fig_hybrid() -> Path:
    img = g.canvas(1400, 560)
    d0 = ImageDraw.Draw(img)
    g.fig_label(
        d0, "Figure  ·  Hybrid Search", "Keyword and meaning, fused, then scoped."
    )
    g.plate(img, (S(40), S(108), S(480), S(420)), 14)
    d = ImageDraw.Draw(img)
    g.draw_rr(d, (S(40), S(108), S(480), S(116)), 0, PRIMARY)
    g.pill(d, (S(60), S(136), S(220), S(168)), PRIMARY_50, "Keyword", PRIMARY, 11)
    g.txt(d, (S(60), S(196)), "Elasticsearch", 20, True, NAVY)
    g.wrapped(
        d,
        S(60),
        S(248),
        "Exact legal and clinical terms. Clause numbers. Party names. The words on the page.",
        S(380),
        14,
        False,
        TEXT,
    )
    g.plate(img, (S(920), S(108), S(1360), S(420)), 14)
    d = ImageDraw.Draw(img)
    g.draw_rr(d, (S(920), S(108), S(1360), S(116)), 0, ORANGE)
    g.pill(d, (S(940), S(136), S(1100), S(168)), ORANGE_50, "Vector", ORANGE, 11)
    g.txt(d, (S(940), S(196)), "Qdrant", 20, True, NAVY)
    g.wrapped(
        d,
        S(940),
        S(248),
        "Meaning. A 384-dimension cosine vector. Nearby ideas when the nouns differ.",
        S(380),
        14,
        False,
        TEXT,
    )
    g.plate(img, (S(500), S(160), S(900), S(380)), 14)
    d = ImageDraw.Draw(img)
    g.pill(
        d, (S(560), S(184), S(840), S(216)), TEAL_50, "Fuse  ·  Equal Weight", TEAL, 11
    )
    g.txt(d, (S(700), S(256)), "One Ranked List", 18, True, NAVY, "mm")
    g.wrapped(
        d,
        S(528),
        S(290),
        "Min-max normalize. Alpha 0.5. Then the scope gate.",
        S(344),
        13,
        False,
        MUTED,
    )
    d.polygon(
        [(S(480), S(250)), (S(500), S(260)), (S(480), S(270))], fill=PRIMARY + (180,)
    )
    d.polygon(
        [(S(920), S(250)), (S(900), S(260)), (S(920), S(270))], fill=ORANGE + (180,)
    )
    d = ImageDraw.Draw(img)
    g.txt(
        d,
        (S(48), S(456)),
        "Shared by chat, agent search, and MCP. If both engines are down, a scoped database lookup. Results are never returned unscoped.",
        13,
        False,
        MUTED,
    )
    g.txt(
        d,
        (S(48), S(496)),
        "§8.2 and “late delivery liability” resolve to the same ranked list.",
        13,
        False,
        MUTED,
    )
    return save(img, "fig-hybrid.png")


def fig_compliance() -> Path:
    img = g.canvas(1400, 580)
    d0 = ImageDraw.Draw(img)
    g.fig_label(d0, "Figure  ·  Compliance", "Controls the agent inherits.")
    items = [
        (40, "Identity", "JWT. TOTP MFA. Session bound to a person.", PRIMARY),
        (
            390,
            "Authorization",
            "RBAC role. ABAC classification. Tenant isolation.",
            ORANGE,
        ),
        (740, "Standards", "HIPAA. PCI-DSS. GDPR-ready. PHI scan and redaction.", TEAL),
        (1090, "Evidence", "Full audit. SIEM export. Secret scanning in CI.", PURPLE),
        (40, "Data", "Field encryption. DLP on export. Watermarks.", GREEN),
        (
            390,
            "Model Path",
            "Ollama first. Air-gap capable. Cloud only as fallback.",
            PRIMARY,
        ),
        (
            740,
            "Tool Scope",
            "Every tool call scoped. Not available — not the text.",
            ORANGE,
        ),
        (
            1090,
            "Step Limits",
            "Six steps by default. Twelve hard cap. Two deep dives.",
            TEAL,
        ),
    ]
    for i, (x, title, body, accent) in enumerate(items):
        y = S(108) if i < 4 else S(330)
        g.plate(img, (S(x), y, S(x + 330), y + S(196)), 12)
        d = ImageDraw.Draw(img)
        g.draw_rr(d, (S(x), y, S(x + 8), y + S(196)), 0, accent)
        g.txt(d, (S(x + 24), y + S(24)), title, 16, True, NAVY)
        g.wrapped(d, S(x + 24), y + S(70), body, S(282), 13, False, TEXT)
    return save(img, "fig-compliance.png")


def fig_promises() -> Path:
    img = g.canvas(1400, 520)
    d0 = ImageDraw.Draw(img)
    g.fig_label(
        d0, "Figure 6  ·  Operating Constraints", "What must hold in regulated work."
    )
    cols = [
        (
            40,
            PRIMARY,
            PRIMARY_50,
            "Scope",
            "Every tool call is filtered to the user's permitted documents. The agent cannot cross a permission boundary.",
        ),
        (
            390,
            ORANGE,
            ORANGE_50,
            "Audit",
            "Think, act, and observe stream live. The full trace returns with the answer.",
        ),
        (
            740,
            TEAL,
            TEAL_50,
            "Private",
            "Your infrastructure. Ollama first. OpenAI only as fallback. Data does not have to leave.",
        ),
        (
            1090,
            PURPLE,
            PURPLE_50,
            "Bounded",
            "A step budget and a repeat-action guard. The run terminates and synthesizes. It does not loop without limit.",
        ),
    ]
    for x, accent, wash, title, body in cols:
        g.plate(img, (S(x), S(108), S(x + 330), S(480)), 14)
        d = ImageDraw.Draw(img)
        g.draw_rr(d, (S(x), S(108), S(x + 330), S(116)), 0, accent)
        g.pill(d, (S(x + 16), S(136), S(x + 140), S(168)), wash, title, accent, 12)
        g.wrapped(d, S(x + 16), S(196), body, S(298), 13.5, False, TEXT, 1.32)
    return save(img, "fig-promises.png")


def fig_who() -> Path:
    img = g.canvas(1400, 480)
    d0 = ImageDraw.Draw(img)
    g.fig_label(d0, "Figure 7  ·  Industries", "Healthcare, legal, and finance.")
    cards = [
        (
            40,
            "Healthcare",
            "HIPAA modes. PHI scan and redaction. Classification on every file.",
            PRIMARY,
        ),
        (
            500,
            "Legal",
            "Cited passages. Scoped libraries. A trace counsel can open.",
            ORANGE,
        ),
        (960, "Finance", "PCI-DSS modes. Dual index. No walk of the book.", TEAL),
    ]
    for x, title, body, accent in cards:
        g.plate(img, (S(x), S(108), S(x + 420), S(320)), 14)
        d = ImageDraw.Draw(img)
        g.draw_rr(d, (S(x), S(108), S(x + 420), S(116)), 0, accent)
        g.txt(d, (S(x + 20), S(148)), title, 18, True, NAVY)
        g.wrapped(d, S(x + 20), S(196), body, S(380), 14, False, TEXT)
    d = ImageDraw.Draw(img)
    g.txt(
        d,
        (S(48), S(360)),
        "JWT  ·  TOTP MFA  ·  RBAC / ABAC  ·  DLP  ·  Field encryption  ·  SIEM export  ·  Secret scanning in CI",
        14,
        True,
        NAVY,
    )
    g.txt(
        d,
        (S(48), S(404)),
        "Stack: FastAPI  ·  Python 3.11  ·  React 18  ·  Postgres  ·  Redis  ·  Celery  ·  Elasticsearch  ·  Qdrant  ·  Docker",
        13,
        False,
        MUTED,
    )
    return save(img, "fig-who.png")


def header_footer(section):
    g.header_footer(section)
    h = section.header.paragraphs[0]
    h.clear()
    r = h.add_run("inDoc")
    g.font_run(r, "Calibri", 9, True, C_PRIMARY)
    r = h.add_run("    Executive Introduction")
    g.font_run(r, "Calibri", 9, False, C_MUTED)
    r = h.add_run(
        "                                                         Confidential"
    )
    g.font_run(r, "Calibri", 9, False, C_MUTED)
    f = section.footer.paragraphs[0]
    f.clear()
    r = f.add_run("Introducing inDoc  ·  Autonomous, But Accountable")
    g.font_run(r, "Calibri", 8.5, False, C_MUTED)
    r = f.add_run("          ")
    r2 = f.add_run()
    g.font_run(r2, "Calibri", 8.5, True, C_PRIMARY)
    g.add_field(r2, "PAGE")
    r3 = f.add_run("  /  ")
    g.font_run(r3, "Calibri", 8.5, False, C_MUTED)
    r4 = f.add_run()
    g.font_run(r4, "Calibri", 8.5, True, C_PRIMARY)
    g.add_field(r4, "NUMPAGES")
    f0 = section.first_page_footer.paragraphs[0]
    f0.clear()
    r = f0.add_run("Shared Oxygen  ·  inDoc Executive Introduction  ·  August 2026")
    g.font_run(r, "Calibri", 8.5, False, C_MUTED)


def build(figs: dict[str, Path]) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = C_TEXT
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.7)
    sec.header_distance = Inches(0.38)
    sec.footer_distance = Inches(0.36)
    header_footer(sec)

    g.para(
        doc,
        "Executive Introduction    ·    August 2026    ·    Shared Oxygen",
        size=9,
        bold=True,
        color=C_PRIMARY,
        after=6,
        keep_next=True,
    )
    g.para(
        doc,
        "Introducing inDoc",
        size=26,
        bold=True,
        color=C_NAVY,
        after=6,
        keep_next=True,
    )
    g.para(
        doc,
        "Autonomous document intelligence on your own infrastructure.",
        size=13,
        italic=True,
        color=C_PRIMARY,
        after=10,
    )

    g.figure(
        doc,
        figs["cover"],
        "inDoc investigates a document library. It does not chat at a file.",
    )

    g.para(
        doc,
        "Document chat products retrieve once and generate an answer. The model does not decide what to do next. That is useful retrieval. It is not an investigation, and it is not a control system for regulated work.",
    )
    g.para(
        doc,
        "inDoc is a private document platform built around three capabilities. An agent investigates: it plans, searches, reads, and re-plans, and it returns a trace a reviewer can open. Hybrid search finds both the exact term and the nearby idea — Elasticsearch for the words, Qdrant for meaning — then applies scope before any result leaves. Compliance binds the run: identity, classification, HIPAA and PCI modes, audit, SIEM, and a local model path. The agent cannot read a file the requesting user is not authorized to see.",
    )
    g.figure(doc, figs["pillars"], "Agentic AI, hybrid search, and compliance.")
    g.quote(doc, "Autonomous, but accountable.")

    g.heading(doc, "01", "The Problem")
    g.para(
        doc,
        "Healthcare, legal, and finance already run large document libraries. The gap is asking a hard question and seeing how the system worked — without giving a model unrestricted access to every file. Retrieve-and-stuff products hand the model a context window. When the answer is wrong, there is no investigation to inspect. When the user is not cleared for a file, policy text is not a control.",
    )
    g.figure(doc, figs["rag"], "Figure 1. RAG is a pipeline. inDoc is a decision loop.")
    g.para(
        doc,
        "inDoc exists so a regulated team can run autonomous research under the same access rules as the rest of the firm. Every tool call is scoped. Every step is recorded. The run is bounded. The model can remain on-premises.",
    )

    g.heading(doc, "02", "A Live Investigation")
    g.para(
        doc,
        "Ask what Q3 revenue was, and its growth rate. The agent does not load the library into a prompt. It searches, identifies the Q3 financial report, reads the file, and returns a cited figure: Q3 2025 revenue was $4.2 million, up 18% year over year.",
    )
    g.figure(
        doc,
        figs["theater"],
        "Think, act, observe — the live instrumentation of an investigation.",
    )
    g.figure(
        doc, figs["trace"], "Figure 2. A product trace. The model chose each move."
    )
    g.para(
        doc,
        "Longer questions continue within a step budget — compare when the question is a difference, summarize when only the gist is required. The run streams as it happens. The public console at sharedoxygen.github.io/indoc replays this loop. The product uses the same endpoints: POST /api/v1/agent/run and POST /api/v1/agent/stream.",
    )

    g.heading(doc, "03", "The Platform")
    g.para(
        doc,
        "A file is virus-scanned, extracted, embedded, and dual-indexed. People chat over a selected set with citations and memory. The agent investigates across what the user is allowed to see. Hybrid search is the shared retrieve path for chat, the agent, and tools. Every hop inherits identity, classification, and tenant. Ollama is first. OpenAI is fallback. Postgres remains the record.",
    )
    g.figure(
        doc,
        figs["loop"],
        "Figure 3. A goal in. A plan. Tools inside scope. An answer and a trace out.",
    )
    g.table2(
        doc,
        ["Capability", "What Ships"],
        [
            [
                "Agentic AI",
                "ReAct loop. Six scope-enforced tools. Live SSE trace. Step budget. Repeat-action guard.",
            ],
            [
                "Hybrid Search",
                "Elasticsearch and Qdrant. Scores fused. Scope applied before results leave.",
            ],
            [
                "Compliance",
                "HIPAA and PCI modes. RBAC and ABAC. Audit. SIEM. Local model. DLP. Field encryption.",
            ],
            [
                "Chat",
                "Multi-document. Streaming. Cited. Memory-aware. Uses hybrid search under scope.",
            ],
            [
                "Pipeline",
                "Virus scan, OCR, embed, dual index. Celery, off the request path.",
            ],
        ],
    )

    g.heading(doc, "04", "Six Tools, One Boundary")
    g.para(
        doc,
        "The agent has six tools. List shows what is in scope. Search is the primary move. Read and summarize are limited to two deep dives per run. Compare is for what changed. Finish returns the answer with citations. Every call uses the same RBAC and ABAC scope as the rest of the platform. A document outside the user's effective set returns “not available,” not the text.",
    )
    g.figure(
        doc,
        figs["tools"],
        "Figure 4. The tools the planner sees. The gate the platform enforces.",
    )

    g.heading(doc, "05", "From File To Ready")
    g.para(
        doc,
        "A file is scanned, extracted — including OCR on images — embedded as a 384-dimension cosine vector, and written to both search engines. Celery performs that work off the request path. Supported formats include PDF, Word, Excel, PowerPoint, text, HTML, mail, and images. Folders and document sets keep the library navigable. If both engines are unavailable, search falls back to a scoped database lookup. Owner, classification, and tenant remain in Postgres.",
    )
    g.figure(
        doc,
        figs["pipe"],
        "Figure 5. From upload to ready. Dual index. Scoped from the first query.",
    )

    g.heading(doc, "06", "Hybrid Search")
    g.para(
        doc,
        "Elasticsearch matches legal and clinical terms — clause numbers, party names, the words on the page. Qdrant matches meaning, even when the nouns differ. Scores are min-max normalized and fused at equal weight. Role, classification, and any document selection are applied before results leave. Chat, the agent's search tool, and MCP share this path. An unscoped result is not returned. §8.2 and “late delivery liability” resolve to the same ranked list.",
    )
    g.figure(
        doc,
        figs["hybrid"],
        "Figure. Hybrid search. Keyword and vector, fused, then scoped.",
    )
    g.figure(doc, figs["scope"], "Allowed documents are visible. The rest are not.")
    g.figure(
        doc, figs["promises"], "Figure 6. Scope, audit, privacy, and a bounded run."
    )

    g.heading(doc, "07", "Compliance")
    g.para(
        doc,
        "Identity, role, classification, and tenant isolation sit in front of every hop. HIPAA and PCI-DSS modes, PHI scan and redaction, DLP on export, watermarks, field encryption, SIEM export, and secret scanning in CI are how the platform is operated. The agent inherits those controls. Healthcare requires that for PHI. Legal requires a trace counsel can open. Finance requires an agent that cannot walk the book.",
    )
    g.figure(
        doc, figs["compliance"], "Figure. Eight controls. The agent inherits each one."
    )
    g.figure(
        doc,
        figs["who"],
        "Figure 7. Healthcare, legal, finance, and the stack beneath them.",
    )

    g.heading(doc, "08", "Deployment")
    g.para(
        doc,
        "inDoc is self-hosted and can run air-gapped. Postgres and Redis are the data plane. Docker runs Elasticsearch, Qdrant, Celery, the API, and operations services. The React application sits on your network. Source and CI are at github.com/sharedoxygen/indoc. The interactive demonstration is at sharedoxygen.github.io/indoc.",
    )
    g.table2(
        doc,
        ["Resource", "Location"],
        [
            ["Demonstration", "https://sharedoxygen.github.io/indoc/"],
            ["Source And CI", "https://github.com/sharedoxygen/indoc"],
            ["Local Application", "http://localhost:5193"],
            ["API", "http://localhost:8001/api/v1/docs"],
        ],
    )

    g.heading(doc, "09", "Position")
    g.para(
        doc,
        "inDoc is an investigation system, not a chat overlay. A goal is planned, executed through tools, and closed with evidence. Hybrid search is the retrieve path. Compliance is the boundary. The trace is complete. Generation can stay on your infrastructure.",
    )
    g.quote(doc, "Autonomous, but accountable. Shared Oxygen, LLC.")
    g.para(
        doc, "© 2026 Shared Oxygen, LLC.", size=10, bold=True, color=C_PRIMARY, after=0
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


def render() -> dict[str, Path]:
    ASSET.mkdir(parents=True, exist_ok=True)
    return {
        "cover": ASSET / "intro-cover-compass.png",
        "theater": ASSET / "intro-think-act-observe.png",
        "scope": ASSET / "intro-scope-wall.png",
        "pillars": fig_three_pillars(),
        "hybrid": fig_hybrid(),
        "compliance": fig_compliance(),
        "rag": fig_rag_vs_agent(),
        "trace": fig_q3_trace(),
        "loop": fig_loop(),
        "tools": fig_tools(),
        "pipe": fig_pipeline(),
        "promises": fig_promises(),
        "who": fig_who(),
    }


if __name__ == "__main__":
    figs = render()
    path = build(figs)
    print(path)
    print("bytes", path.stat().st_size)
