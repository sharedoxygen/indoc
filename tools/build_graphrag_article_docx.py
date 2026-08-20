#!/usr/bin/env python3
# flake8: noqa: E501
"""Executive inDoc GraphRAG briefing — original precision figures, tight Word layout."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from PIL import Image, ImageDraw, ImageFilter, ImageFont

ROOT = Path(__file__).resolve().parents[1]
ASSET = ROOT / "private-docs/gtm/assets/graphrag-article"
OUT = (
    ROOT
    / "private-docs/gtm/how-indoc-answers-questions-hidden-across-a-document-library.docx"
)

CANVAS = (247, 249, 252)
PAPER = (255, 255, 255)
PRIMARY = (25, 118, 210)
PRIMARY_50 = (232, 244, 253)
PRIMARY_100 = (209, 233, 251)
PRIMARY_DARK = (13, 71, 161)
PRIMARY_600 = (21, 101, 192)
ORANGE = (230, 126, 0)
ORANGE_50 = (255, 243, 224)
TEAL = (0, 121, 107)
TEAL_50 = (224, 242, 241)
PURPLE = (123, 31, 162)
PURPLE_50 = (243, 229, 245)
GREEN = (46, 125, 50)
GREEN_50 = (232, 245, 233)
RED = (198, 40, 40)
RED_50 = (255, 235, 238)
TEXT = (33, 37, 41)
MUTED = (90, 98, 108)
LINE = (221, 226, 232)
NAVY = (15, 55, 120)
INK = (22, 28, 36)

FONT_REG = "/Library/Fonts/Microsoft/Arial.ttf"
FONT_BOLD = "/Library/Fonts/Microsoft/Arial Bold.ttf"
FONT_ITA = "/Library/Fonts/Microsoft/Arial Italic.ttf"
SCALE = 4

_PRESERVE = {
    "indoc": "inDoc",
    "graphrag": "GraphRAG",
    "elasticsearch": "Elasticsearch",
    "postgresql": "PostgreSQL",
    "qdrant": "Qdrant",
    "ollama": "Ollama",
    "react": "ReAct",
    "rag": "RAG",
    "nlp": "NLP",
    "ocr": "OCR",
    "rbac": "RBAC",
    "abac": "ABAC",
    "mcp": "MCP",
    "msa": "MSA",
    "dpa": "DPA",
    "sow": "SOW",
    "eu": "EU",
    "q3": "Q3",
    "api": "API",
    "llm": "LLM",
    "ai": "AI",
    "hipaa": "HIPAA",
    "pci": "PCI",
    "dss": "DSS",
    "gdpr": "GDPR",
    "siem": "SIEM",
    "jwt": "JWT",
    "totp": "TOTP",
    "mfa": "MFA",
    "dlp": "DLP",
    "phi": "PHI",
    "sse": "SSE",
    "openai": "OpenAI",
}


def init_cap(text: str) -> str:
    """Title-case labels. Preserve product names and identifiers."""
    if not text:
        return text
    out = []
    for i, raw in enumerate(text.split(" ")):
        if raw in {"", "·", "—", "–", "/", "→"}:
            out.append(raw)
            continue
        lead = ""
        core = raw
        trail = ""
        while core and core[0] in "§·—–-/":
            lead += core[0]
            core = core[1:]
        while core and core[-1] in ".,:;·—–-/":
            trail = core[-1] + trail
            core = core[:-1]
        if not core:
            out.append(raw)
            continue
        if any(ch.isdigit() for ch in core) and sum(ch.isdigit() for ch in core) >= 3:
            out.append(raw)
            continue
        key = core.lower()
        if key in _PRESERVE:
            out.append(f"{lead}{_PRESERVE[key]}{trail}")
        elif core[:1].islower() and any(ch.isupper() for ch in core[1:]):
            out.append(raw)
        elif "-" in core:
            bits = []
            for part in core.split("-"):
                lk = part.lower()
                if lk in _PRESERVE:
                    bits.append(_PRESERVE[lk])
                elif part:
                    bits.append(part[:1].upper() + part[1:].lower())
                else:
                    bits.append(part)
            out.append(f"{lead}{'-'.join(bits)}{trail}")
        else:
            out.append(f"{lead}{core[:1].upper()}{core[1:].lower()}{trail}")
    return " ".join(out)


C_PRIMARY = RGBColor(25, 118, 210)
C_NAVY = RGBColor(15, 55, 120)
C_TEXT = RGBColor(33, 37, 41)
C_MUTED = RGBColor(90, 98, 108)
C_INK = RGBColor(22, 28, 36)


def F(size: int, bold: bool = False, italic: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD if bold else (FONT_ITA if italic else FONT_REG)
    return ImageFont.truetype(path, int(size * SCALE))


def canvas(w: int, h: int) -> Image.Image:
    return Image.new("RGBA", (w * SCALE, h * SCALE), CANVAS + (255,))


def S(v: float) -> int:
    return int(v * SCALE)


def draw_rr(d: ImageDraw.ImageDraw, box, r, fill, outline=None, width=1.25):
    d.rounded_rectangle(
        box,
        radius=S(r),
        fill=fill if fill is None or len(fill) == 4 else fill + (255,),
        outline=None if outline is None else outline + (255,),
        width=S(width),
    )


def plate(img: Image.Image, box, r=14, fill=PAPER, outline=LINE, shadow=True):
    x0, y0, x1, y1 = map(int, box)
    if shadow:
        layer = Image.new("RGBA", img.size, (0, 0, 0, 0))
        ld = ImageDraw.Draw(layer)
        ld.rounded_rectangle(
            (x0 + S(3), y0 + S(6), x1 + S(3), y1 + S(8)),
            radius=S(r),
            fill=(15, 23, 42, 22),
        )
        img.alpha_composite(layer.filter(ImageFilter.GaussianBlur(S(7))))
    d = ImageDraw.Draw(img)
    draw_rr(d, (x0, y0, x1, y1), r, fill, outline, 1.15)
    return d


def txt(d, xy, s, size=13, bold=False, fill=TEXT, anchor="lt", italic=False):
    d.text(xy, s, font=F(size, bold, italic), fill=fill + (255,), anchor=anchor)


def measure(d, s, size=13, bold=False, italic=False) -> float:
    return d.textlength(s, font=F(size, bold, italic))


def wrap(d, s, max_w, size=13, bold=False, italic=False) -> list[str]:
    words = s.split()
    lines, cur = [], ""
    for w in words:
        trial = (cur + " " + w).strip()
        if measure(d, trial, size, bold, italic) <= max_w:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines or [""]


def wrapped(
    d,
    x,
    y,
    s,
    max_w,
    size=13,
    bold=False,
    fill=TEXT,
    leading=1.28,
    italic=False,
    anchor_x="lt",
):
    lines = wrap(d, s, max_w, size, bold, italic)
    fh = size * leading
    for i, line in enumerate(lines):
        txt(d, (x, y + S(fh * i)), line, size, bold, fill, anchor_x, italic)
    return len(lines) * fh


def pill(d, box, fill, label, fg=PAPER, size=10.5):
    draw_rr(d, box, 999, fill)
    x0, y0, x1, y1 = box
    txt(d, ((x0 + x1) / 2, (y0 + y1) / 2), label, size, True, fg, "mm")


def save(img: Image.Image, name: str) -> Path:
    rgb = Image.new("RGB", img.size, CANVAS)
    rgb.paste(img, mask=img.split()[-1])
    w = 2400
    h = max(1, int(rgb.height * w / rgb.width))
    rgb = rgb.resize((w, h), Image.Resampling.LANCZOS)
    out = ASSET / name
    rgb.save(out, "PNG", optimize=True, dpi=(300, 300))
    return out


def fig_label(d, kicker, title, x=48, y=22):
    txt(d, (S(x), S(y)), init_cap(kicker), 10.5, True, PRIMARY_600)
    wrapped(d, S(x), S(y + 22), init_cap(title), S(1280), 17.5, True, NAVY, 1.2)


def facsimile(
    d,
    box,
    title,
    meta,
    body_lines,
    highlight=None,
    accent=PRIMARY,
    footer="Confidential  ·  Page 4 of 28",
):
    x0, y0, x1, y1 = box
    draw_rr(d, box, 8, PAPER, LINE, 1)
    draw_rr(d, (x0, y0, x1, y0 + S(36)), 0, accent)
    # square the header bottom
    d.rectangle((x0, y0 + S(20), x1, y0 + S(36)), fill=accent + (255,))
    txt(d, (x0 + S(14), y0 + S(18)), init_cap(title), 11, True, PAPER, "lm")
    txt(d, (x0 + S(14), y0 + S(50)), meta, 9.5, False, MUTED, "lt")
    d.line(
        (x0 + S(14), y0 + S(68), x1 - S(14), y0 + S(68)), fill=LINE + (255,), width=S(1)
    )
    top = y0 + S(82)
    for i, line in enumerate(body_lines):
        ly = top + S(i * 18)
        color = accent if highlight is not None and i == highlight else (210, 216, 222)
        if highlight is not None and i == highlight:
            d.rectangle(
                (x0 + S(12), ly - S(3), x1 - S(12), ly + S(16)),
                fill=PRIMARY_50 + (255,),
            )
        d.rounded_rectangle(
            (x0 + S(14), ly, x0 + S(14) + S(line), ly + S(9)),
            radius=S(3),
            fill=color + (255,),
        )
    txt(d, (x0 + S(14), y1 - S(16)), footer, 8, False, MUTED, "lm")


def fig_01_case() -> Path:
    img = canvas(1400, 720)
    d0 = ImageDraw.Draw(img)
    fig_label(
        d0,
        "Figure 1  ·  The case",
        "One clause has a home. The pattern lives across the portfolio.",
    )
    d = plate(img, (S(40), S(88), S(676), S(688)), 16)
    txt(d, (S(64), S(112)), "Located Question", 10, True, PRIMARY)
    txt(
        d,
        (S(64), S(134)),
        "Which clause assigns late-delivery liability?",
        14.5,
        True,
        NAVY,
    )
    facsimile(
        d,
        (S(64), S(172), S(652), S(548)),
        "Master Services Agreement  —  Acme Corp",
        "Effective 12 Mar 2024   ·   Counterparty: Acme Holdings   ·   Classification: Restricted",
        [210, 248, 188, 260, 240, 160, 252, 200, 228, 140, 190, 170],
        highlight=6,
        accent=PRIMARY,
        footer="inDoc  ·  doc 9f6c…   ·   §8.2 Late Delivery   ·   Restricted",
    )
    txt(
        d,
        (S(64), S(572)),
        "§8.2  Late Delivery. Vendor bears liability for delay",
        12,
        False,
        TEXT,
    )
    txt(
        d,
        (S(64), S(594)),
        "beyond the committed date, capped at six months of fees.",
        12,
        False,
        TEXT,
    )
    txt(
        d,
        (S(64), S(632)),
        "inDoc today: hybrid search finds the passage. The agent cites it.",
        11.5,
        False,
        MUTED,
    )
    d = plate(img, (S(700), S(88), S(1360), S(688)), 16)
    txt(d, (S(724), S(112)), "Corpus Question", 10, True, ORANGE)
    txt(
        d,
        (S(724), S(134)),
        "Which liability patterns recur since 2021?",
        14.5,
        True,
        NAVY,
    )
    vendors = [
        ("Acme MSA", True),
        ("Northline MSA", True),
        ("Helix DPA", False),
        ("Vesper MSA", True),
        ("Oak & Pine", False),
        ("Meridian SOW", False),
        ("Brightlane MSA", True),
        ("Cove Logistics", False),
        ("Pinnacle MSA", True),
        ("Harbor QMSA", False),
        ("Lumen Addendum", True),
        ("Arcadia MSA", False),
    ]
    for i, (name, hit) in enumerate(vendors):
        col, row = i % 3, i // 3
        x = S(724) + col * S(200)
        y = S(176) + row * S(96)
        fill = ORANGE_50 if hit else (245, 247, 250)
        edge = ORANGE if hit else LINE
        draw_rr(d, (x, y, x + S(184), y + S(82)), 8, fill, edge, 1.2)
        txt(d, (x + S(12), y + S(18)), name, 11.5, True, NAVY if hit else MUTED)
        txt(
            d,
            (x + S(12), y + S(42)),
            "6-Mo Fee Cap" if hit else "Standard Form",
            10,
            False,
            ORANGE if hit else MUTED,
        )
        txt(
            d,
            (x + S(12), y + S(60)),
            "Shared Deviation" if hit else "No Match",
            9.5,
            False,
            ORANGE if hit else (160, 166, 174),
        )
    txt(
        d,
        (S(724), S(632)),
        "6 of 12 MSAs share the same cap. No single file contains that fact.",
        11.5,
        False,
        MUTED,
    )
    return save(img, "fig-01-case.png")


def fig_02_stack() -> Path:
    img = canvas(1400, 620)
    d0 = ImageDraw.Draw(img)
    fig_label(
        d0,
        "Figure 2  ·  What ships today",
        "Every hop is scoped. The agent is budgeted. The citation is checkable.",
    )
    stages = [
        (
            40,
            "Ingest",
            "Virus scan → extract → OCR → dual index",
            "Celery, off the request path",
            PRIMARY,
            PRIMARY_50,
        ),
        (
            385,
            "Retrieve",
            "Keyword + vector, fused equally",
            "Elasticsearch · Qdrant 384-d",
            ORANGE,
            ORANGE_50,
        ),
        (
            730,
            "Reason",
            "Chat on a selected set, or ReAct",
            "Plan · search · read · finish",
            TEAL,
            TEAL_50,
        ),
        (
            1075,
            "Govern",
            "RBAC · ABAC · selection · audit",
            "No text outside effective scope",
            PURPLE,
            PURPLE_50,
        ),
    ]
    d = ImageDraw.Draw(img)
    for x, title, line1, line2, accent, wash in stages:
        plate(img, (S(x), S(100), S(x + 325), S(340)), 14)
        dd = ImageDraw.Draw(img)
        draw_rr(dd, (S(x), S(100), S(x + 325), S(108)), 0, accent)
        dd.rectangle((S(x), S(106), S(x + 325), S(108)), fill=accent + (255,))
        pill(
            dd,
            (S(x + 16), S(124), S(x + 118), S(150)),
            wash,
            init_cap(title),
            accent,
            10,
        )
        wrapped(dd, S(x + 16), S(172), init_cap(line1), S(292), 14, True, NAVY)
        wrapped(dd, S(x + 16), S(230), init_cap(line2), S(292), 12.5, False, MUTED)
        if x < 1075:
            dd.polygon(
                [(S(x + 333), S(210)), (S(x + 349), S(220)), (S(x + 333), S(230))],
                fill=PRIMARY + (180,),
            )
    stores = [
        (40, "PostgreSQL", "Record, owner, classification"),
        (385, "Elasticsearch", "The words on the page"),
        (730, "Qdrant", "Meaning, 384-dimension cosine"),
        (1075, "Ollama", "Local model, cloud fallback"),
    ]
    for x, title, sub in stores:
        plate(img, (S(x), S(372), S(x + 325), S(500)), 12)
        dd = ImageDraw.Draw(img)
        txt(dd, (S(x + 16), S(400)), init_cap(title), 14, True, NAVY)
        wrapped(dd, S(x + 16), S(430), init_cap(sub), S(292), 12, False, MUTED)
    dd = ImageDraw.Draw(img)
    txt(
        dd,
        (S(48), S(540)),
        "Agent budget: 6 steps default, 12 hard cap, 2 deep dives. Unbounded library walks are refused.",
        12.5,
        False,
        MUTED,
    )
    txt(
        dd,
        (S(48), S(568)),
        "Shared path: chat, agent search_documents, and MCP search all use the same scoped hybrid fuse.",
        12.5,
        False,
        MUTED,
    )
    return save(img, "fig-02-stack.png")


def fig_03_gap() -> Path:
    img = canvas(1400, 680)
    d0 = ImageDraw.Draw(img)
    fig_label(
        d0,
        "Figure 3  ·  The resemblance gap",
        "Nearest neighbors of “recur most often” are vocabulary cousins — not the histogram.",
    )
    plate(img, (S(40), S(96), S(680), S(648)), 16)
    d = ImageDraw.Draw(img)
    txt(d, (S(64), S(120)), "What Similarity Returns", 10, True, PRIMARY)
    txt(
        d,
        (S(64), S(144)),
        "Query: Which Failure Causes Recur Most Often?",
        13,
        True,
        NAVY,
    )
    hits = [
        ("PM-2019-04", "Recurring pager alerts in checkout", 0.81, False),
        ("PM-2021-11", "Frequent timeouts after release", 0.77, False),
        ("RUN-88", "How to mark an incident recurring", 0.74, False),
        ("PM-2020-02", "Again: retry storm on payments", 0.69, False),
        ("GLOSS-7", "Glossary — recurring vs systemic", 0.64, False),
    ]
    for i, (doc_id, title, score, _) in enumerate(hits):
        y = S(184) + i * S(76)
        draw_rr(
            d,
            (S(64), y, S(656), y + S(66)),
            8,
            PRIMARY_50 if i == 0 else (245, 247, 250),
            LINE,
        )
        txt(d, (S(80), y + S(16)), doc_id, 10, True, PRIMARY)
        txt(d, (S(80), y + S(38)), title, 12.5, False, TEXT)
        txt(d, (S(620), y + S(33)), f"{score:.2f}", 12, True, MUTED, "rm")
    txt(
        d,
        (S(64), S(600)),
        "These files say “recurring.” They do not tally root causes.",
        12,
        False,
        MUTED,
    )

    plate(img, (S(704), S(96), S(1360), S(648)), 16)
    d = ImageDraw.Draw(img)
    txt(d, (S(728), S(120)), "What The Question Wants", 10, True, ORANGE)
    txt(
        d,
        (S(728), S(144)),
        "Distribution across 214 postmortems, 2021–2025",
        13,
        True,
        NAVY,
    )
    bars = [
        ("Timeouts after deploys", 92, ORANGE, "47 files"),
        ("Retry storms", 78, PRIMARY, "39 files"),
        ("Stale configuration", 54, TEAL, "27 files"),
        ("Upstream dependency lag", 41, PURPLE, "21 files"),
        ("Handoff / ownership gap", 28, GREEN, "14 files"),
    ]
    for i, (label, pct, color, n) in enumerate(bars):
        y = S(188) + i * S(78)
        txt(d, (S(728), y), label, 12.5, True, TEXT)
        txt(d, (S(1336), y), n, 11, False, MUTED, "rt")
        draw_rr(d, (S(728), y + S(24), S(1336), y + S(44)), 6, PRIMARY_50)
        draw_rr(d, (S(728), y + S(24), S(728 + 6.08 * pct), y + S(44)), 6, color)
    txt(
        d,
        (S(728), S(600)),
        "A property of the collection. It does not occupy one vector.",
        12,
        False,
        MUTED,
    )
    return save(img, "fig-03-gap.png")


def fig_04_sentence() -> Path:
    img = canvas(1400, 700)
    d0 = ImageDraw.Draw(img)
    fig_label(
        d0,
        "Figure 4  ·  One sentence, mapped",
        "Extraction must point back to the paragraph. No pointer, no audit.",
    )
    plate(img, (S(40), S(96), S(1360), S(250)), 14)
    d = ImageDraw.Draw(img)
    txt(
        d,
        (S(64), S(118)),
        "SOURCE  ·  INC-2024-031  ·  Checkout timeout postmortem  ·  Restricted",
        10,
        True,
        PRIMARY,
    )
    sentence = "The checkout service began returning timeouts after the Payments team deployed the new retry handler on 3 March."
    wrapped(d, S(64), S(148), sentence, S(1240), 18, False, INK, 1.35)
    txt(
        d,
        (S(64), S(214)),
        "Text unit T-1844   ·   page 3, lines 12–14   ·   this pointer travels with every node below",
        11.5,
        False,
        MUTED,
    )

    entities = [
        (
            40,
            "Checkout service",
            "Entity  ·  Service",
            "Owns the failing request path",
            PRIMARY,
            PRIMARY_50,
        ),
        (
            380,
            "Payments team",
            "Entity  ·  Team",
            "Shipped the change",
            ORANGE,
            ORANGE_50,
        ),
        (
            720,
            "Retry handler",
            "Entity  ·  Component",
            "New logic on 3 March",
            TEAL,
            TEAL_50,
        ),
        (
            1060,
            "3 March deploy",
            "Claim  ·  Time-bound",
            "Deploy preceded timeouts",
            PURPLE,
            PURPLE_50,
        ),
    ]
    for x, title, kind, desc, accent, wash in entities:
        plate(img, (S(x), S(278), S(x + 320), S(430)), 12)
        dd = ImageDraw.Draw(img)
        pill(
            dd,
            (S(x + 14), S(296), S(x + 200), S(320)),
            wash,
            init_cap(kind),
            accent,
            8.5,
        )
        txt(dd, (S(x + 16), S(340)), init_cap(title), 14, True, NAVY)
        wrapped(dd, S(x + 16), S(370), desc, S(288), 12, False, MUTED)

    plate(img, (S(40), S(454), S(1360), S(668)), 14)
    d = ImageDraw.Draw(img)
    txt(d, (S(64), S(478)), "Relationships Retained From T-1844", 10, True, PRIMARY)
    rels = [
        ("Payments team", "deployed", "Retry handler", "3 March"),
        ("Retry handler", "preceded", "Checkout timeouts", "causal order"),
        ("Checkout service", "exhibited", "Timeouts", "symptom"),
    ]
    for i, (src, rel, dst, note) in enumerate(rels):
        y = S(514) + i * S(46)
        draw_rr(d, (S(64), y, S(1316), y + S(38)), 8, (245, 247, 250), LINE)
        txt(d, (S(84), y + S(19)), src, 12.5, True, NAVY, "lm")
        pill(d, (S(430), y + S(8), S(600), y + S(30)), PRIMARY_50, rel, PRIMARY, 10)
        txt(d, (S(628), y + S(19)), dst, 12.5, True, NAVY, "lm")
        txt(d, (S(1288), y + S(19)), note, 11, False, MUTED, "rm")
    return save(img, "fig-04-sentence.png")


def fig_05_path() -> Path:
    img = canvas(1400, 640)
    d0 = ImageDraw.Draw(img)
    fig_label(
        d0,
        "Figure 5  ·  A path no single file contains",
        "Similarity returns lookalikes. The graph walks the connection.",
    )
    docs = [
        (
            48,
            "ADR-014  ·  Design",
            "Maya Chen named as owner",
            "of the payments retry path.",
            PRIMARY,
            "Engineer",
        ),
        (
            508,
            "RB-221  ·  Runbook",
            "Payments retry path is",
            "implemented in checkout-api.",
            TEAL,
            "Service",
        ),
        (
            968,
            "INC-2024-031  ·  Incident",
            "checkout-api timeouts after",
            "the 3 March retry deploy.",
            ORANGE,
            "Incident",
        ),
    ]
    for x, title, l1, l2, accent, node in docs:
        plate(img, (S(x), S(108), S(x + 384), S(340)), 14)
        d = ImageDraw.Draw(img)
        draw_rr(d, (S(x), S(108), S(x + 8), S(340)), 0, accent)
        txt(d, (S(x + 28), S(132)), title, 12, True, accent)
        wrapped(d, S(x + 28), S(168), l1 + " " + l2, S(332), 14.5, False, INK, 1.3)
        pill(
            d,
            (S(x + 28), S(280), S(x + 150), S(308)),
            accent,
            init_cap(node),
            PAPER,
            10,
        )
    d = ImageDraw.Draw(img)
    # connectors
    for x1, x2 in ((432, 508), (892, 968)):
        y = S(224)
        d.line((S(x1), y, S(x2), y), fill=PRIMARY + (255,), width=S(2.4))
        d.polygon(
            [(S(x2 - 10), y - S(7)), (S(x2), y), (S(x2 - 10), y + S(7))],
            fill=PRIMARY + (255,),
        )
    plate(img, (S(48), S(368), S(1352), S(608)), 14)
    d = ImageDraw.Draw(img)
    txt(
        d,
        (S(72), S(392)),
        "The Walk  ·  Maya Chen  →  Payments Retry Path  →  checkout-api  →  INC-2024-031",
        12,
        True,
        NAVY,
    )
    steps = [
        ("01", "Design names the engineer"),
        ("02", "Runbook names the service"),
        ("03", "Incident names the outage"),
        ("04", "No document names all three"),
    ]
    for i, (n, label) in enumerate(steps):
        x = S(72) + i * S(318)
        draw_rr(
            d, (x, S(436), x + S(298), S(556)), 10, PRIMARY_50 if i < 3 else ORANGE_50
        )
        txt(d, (x + S(16), S(460)), n, 12, True, PRIMARY if i < 3 else ORANGE)
        wrapped(d, x + S(16), S(490), label, S(266), 13.5, True, NAVY)
    return save(img, "fig-05-path.png")


def fig_06_operating() -> Path:
    img = canvas(1400, 760)
    d0 = ImageDraw.Draw(img)
    fig_label(
        d0,
        "Figure 6  ·  Operating model",
        "Build the survey before the question. Route the question to the right index.",
    )
    steps = [
        ("1", "Slice", "Text units", PRIMARY),
        ("2", "Extract", "Entities, links", ORANGE),
        ("3", "Merge", "~75% of cost", RED),
        ("4", "Claim", "Time-bounded", TEAL),
        ("5", "Cluster", "Hierarchy", PURPLE),
        ("6", "Report", "Community briefs", GREEN),
    ]
    for i, (n, title, sub, accent) in enumerate(steps):
        x = 40 + i * 226
        plate(img, (S(x), S(100), S(x + 210), S(268)), 12)
        d = ImageDraw.Draw(img)
        draw_rr(d, (S(x), S(100), S(x + 210), S(106)), 0, accent)
        txt(d, (S(x + 16), S(128)), n, 12, True, accent)
        txt(d, (S(x + 16), S(160)), init_cap(title), 16, True, NAVY)
        txt(d, (S(x + 16), S(198)), init_cap(sub), 12, False, MUTED)
        if i < 5:
            d.polygon(
                [(S(x + 214), S(176)), (S(x + 226), S(184)), (S(x + 214), S(192))],
                fill=PRIMARY + (170,),
            )
    d = ImageDraw.Draw(img)
    txt(
        d,
        (S(48), S(288)),
        "Community reports inherit the highest classification of their sources. Tenants never share nodes.",
        12,
        False,
        MUTED,
    )

    modes = [
        (
            40,
            "Hybrid search",
            "inDoc today",
            "The question looks like a passage. A file or clause is named.",
            PRIMARY,
            "Located",
        ),
        (
            500,
            "Local graph",
            "Gather and rank",
            "An entity is named. Expand neighbors, claims, and source units.",
            TEAL,
            "Connection",
        ),
        (
            960,
            "Global graph",
            "Map, then reduce",
            "The question asks what recurs, differs, or dominates.",
            ORANGE,
            "Pattern",
        ),
    ]
    for x, title, chip, body, accent, kind in modes:
        plate(img, (S(x), S(328), S(x + 440), S(560)), 14)
        d = ImageDraw.Draw(img)
        draw_rr(d, (S(x), S(328), S(x + 440), S(336)), 0, accent)
        pill(
            d,
            (S(x + 16), S(352), S(x + 168), S(378)),
            accent,
            init_cap(chip),
            PAPER,
            9.5,
        )
        txt(d, (S(x + 16), S(400)), init_cap(title), 16, True, NAVY)
        txt(d, (S(x + 16), S(430)), init_cap(kind), 11, True, accent)
        wrapped(d, S(x + 16), S(460), body, S(408), 13, False, TEXT)

    costs = [
        (
            40,
            "Hybrid RAG",
            "Index and query both stay inexpensive.",
            0.22,
            0.12,
            PRIMARY,
            TEAL,
            "Index",
            "Query",
        ),
        (
            500,
            "Full GraphRAG",
            "Pay at ingest. Global questions pay again.",
            0.78,
            0.30,
            ORANGE,
            PURPLE,
            "Extract",
            "Global query",
        ),
        (
            960,
            "Lazy graph",
            "Cheap to build. Model spend moves to query.",
            0.16,
            0.48,
            GREEN,
            RED,
            "NLP index",
            "Query model",
        ),
    ]
    for x, title, sub, a, b, c1, c2, l1, l2 in costs:
        plate(img, (S(x), S(584), S(x + 440), S(728)), 12)
        d = ImageDraw.Draw(img)
        txt(d, (S(x + 16), S(604)), init_cap(title), 13.5, True, NAVY)
        wrapped(d, S(x + 16), S(630), sub, S(408), 11, False, MUTED)
        draw_rr(d, (S(x + 16), S(672), S(x + 16 + 400 * a), S(694)), 5, c1)
        draw_rr(
            d,
            (S(x + 16 + 400 * a + 8), S(672), S(x + 24 + 400 * a + 400 * b), S(694)),
            5,
            c2,
        )
        txt(d, (S(x + 16), S(708)), f"{l1}   ·   {l2}", 10, False, MUTED)
    return save(img, "fig-06-operating.png")


def fig_07_decision() -> Path:
    img = canvas(1400, 620)
    d0 = ImageDraw.Draw(img)
    fig_label(
        d0,
        "Figure 7  ·  Product stance",
        "Keep the system that is already safe. Add a map. Do not sell a miracle.",
    )
    cols = [
        (
            40,
            GREEN,
            GREEN_50,
            "Keep",
            "Hybrid search. Selected chat. The ReAct agent. Dual index. Local model. Scope on every hop. These are why a regulated team can run autonomous research at all.",
        ),
        (
            496,
            PRIMARY,
            PRIMARY_50,
            "Add",
            "An extract job on the existing pipeline. A graph store beside Elasticsearch and Qdrant. Community reports as classified documents. Two agent tools: local and global graph search.",
        ),
        (
            952,
            RED,
            RED_50,
            "Do not",
            "Unbounded walks. Cross-tenant nodes. Community reports that wash Restricted text into an Internal answer. The claim that GraphRAG makes the model truthful.",
        ),
    ]
    for x, accent, wash, title, body in cols:
        plate(img, (S(x), S(100), S(x + 432), S(360)), 14)
        d = ImageDraw.Draw(img)
        draw_rr(d, (S(x), S(100), S(x + 432), S(108)), 0, accent)
        pill(
            d,
            (S(x + 16), S(124), S(x + 118), S(152)),
            wash,
            init_cap(title),
            accent,
            11,
        )
        wrapped(d, S(x + 16), S(176), body, S(400), 13.5, False, TEXT, 1.32)
    plate(img, (S(40), S(388), S(1360), S(588)), 14)
    d = ImageDraw.Draw(img)
    txt(d, (S(64), S(412)), "The Honest Last Hop", 10, True, PRIMARY)
    hops = [
        ("Ask", "Corpus question"),
        ("Scope", "Effective documents only"),
        ("Graph", "Rated community points"),
        ("Confirm", "Open the source files"),
        ("Answer", "Citations and trace"),
    ]
    for i, (title, sub) in enumerate(hops):
        x = S(64) + i * S(256)
        draw_rr(
            d, (x, S(448), x + S(236), S(548)), 10, PRIMARY_50 if i != 3 else TEAL_50
        )
        txt(d, (x + S(16), S(472)), f"{i+1:02d}  {init_cap(title)}", 13.5, True, NAVY)
        wrapped(d, x + S(16), S(504), init_cap(sub), S(204), 12, False, MUTED)
        if i < 4:
            d.polygon(
                [(x + S(240), S(492)), (x + S(252), S(498)), (x + S(240), S(504))],
                fill=PRIMARY + (180,),
            )
    return save(img, "fig-07-decision.png")


# ---------- Word ----------


def shade(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **sides):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge, spec in sides.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in spec.items():
            el.set(qn(f"w:{k}"), str(v))
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_table_borders(table, color="D7DEE6", sz="4"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    tr_pr.append(el)


def keep(p, nxt=True):
    p_pr = p._p.get_or_add_pPr()
    if nxt:
        el = OxmlElement("w:keepNext")
        p_pr.append(el)
    kl = OxmlElement("w:keepLines")
    p_pr.append(kl)
    wo = OxmlElement("w:widowControl")
    wo.set(qn("w:val"), "true")
    p_pr.append(wo)


def p_border(p, color="1976D2"):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def font_run(run, name="Calibri", size=11, bold=False, color=C_TEXT, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def para(
    doc,
    text="",
    *,
    size=11,
    bold=False,
    color=C_TEXT,
    after=8,
    before=0,
    align="left",
    italic=False,
    keep_next=False,
):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.widow_control = True
    keep(p, keep_next)
    if text:
        run = p.add_run(text)
        font_run(run, "Calibri", size, bold, color, italic)
    return p


def rich(doc, parts, **kwargs):
    p = para(doc, "", **kwargs)
    for text, opts in parts:
        run = p.add_run(text)
        font_run(
            run,
            "Calibri",
            opts.get("size", kwargs.get("size", 11)),
            opts.get("bold", False),
            opts.get("color", kwargs.get("color", C_TEXT)),
            opts.get("italic", False),
        )
    return p


def heading(doc, number: str, title: str):
    p = para(doc, "", size=14, before=16, after=8, keep_next=True)
    r1 = p.add_run(f"{number}   ")
    font_run(r1, "Calibri", 14, True, C_PRIMARY)
    r2 = p.add_run(init_cap(title))
    font_run(r2, "Calibri", 14, True, C_NAVY)
    p_border(p)
    keep(p, True)


def figure(doc, path: Path, caption: str):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(table)
    for row in table.rows:
        cant_split(row)
    c0 = table.cell(0, 0)
    c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    keep(p, True)
    p.add_run().add_picture(str(path), width=Inches(6.8))
    c1 = table.cell(1, 0)
    shade(c1, "F4F6F8")
    p2 = c1.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Pt(6)
    keep(p2, False)
    run = p2.add_run(init_cap(caption))
    font_run(run, "Calibri", 9, False, C_MUTED, True)
    para(doc, "", size=4, after=6)


def quote(doc, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(table)
    cant_split(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, "E8F4FD")
    set_cell_border(
        cell,
        left={"val": "single", "sz": "20", "color": "1976D2"},
        top={"val": "nil"},
        bottom={"val": "nil"},
        right={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    keep(p, False)
    run = p.add_run(text)
    font_run(run, "Calibri", 11.5, False, C_NAVY, True)
    para(doc, "", size=4, after=8)


def table2(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "D1E9FB", "6")
    for row in table.rows:
        cant_split(row)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shade(cell, "1565C0")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(init_cap(h))
        font_run(run, "Calibri", 10, True, RGBColor(255, 255, 255))
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            shade(cell, "F7F9FC" if r % 2 == 0 else "FFFFFF")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            font_run(run, "Calibri", 10.5, False, C_TEXT)
    para(doc, "", size=4, after=8)


def add_field(run, instr: str):
    r = run._r
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    t = OxmlElement("w:instrText")
    t.set(qn("xml:space"), "preserve")
    t.text = f" {instr} "
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "end")
    r.append(b)
    r.append(t)
    r.append(e)


def header_footer(section):
    section.different_first_page_header_footer = True
    h0 = section.first_page_header.paragraphs[0]
    h0.text = ""
    f0 = section.first_page_footer.paragraphs[0]
    f0.text = ""
    run = f0.add_run("Shared Oxygen  ·  inDoc Executive Briefing  ·  August 2026")
    font_run(run, "Calibri", 8.5, False, C_MUTED)

    h = section.header.paragraphs[0]
    h.clear()
    r = h.add_run("inDoc")
    font_run(r, "Calibri", 9, True, C_PRIMARY)
    r = h.add_run("    Executive Briefing")
    font_run(r, "Calibri", 9, False, C_MUTED)
    r = h.add_run(
        "                                                                   Confidential"
    )
    font_run(r, "Calibri", 9, False, C_MUTED)

    f = section.footer.paragraphs[0]
    f.clear()
    r = f.add_run("Questions Hidden Across A Document Library")
    font_run(r, "Calibri", 8.5, False, C_MUTED)
    r = f.add_run("          ")
    r2 = f.add_run()
    font_run(r2, "Calibri", 8.5, True, C_PRIMARY)
    add_field(r2, "PAGE")
    r3 = f.add_run("  /  ")
    font_run(r3, "Calibri", 8.5, False, C_MUTED)
    r4 = f.add_run()
    font_run(r4, "Calibri", 8.5, True, C_PRIMARY)
    add_field(r4, "NUMPAGES")


def build(figs: dict[str, Path]) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = C_TEXT
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.7)
    sec.header_distance = Inches(0.38)
    sec.footer_distance = Inches(0.36)
    header_footer(sec)

    para(
        doc,
        "Executive Briefing    ·    August 2026    ·    Original inDoc Analysis",
        size=9,
        bold=True,
        color=C_PRIMARY,
        after=6,
        keep_next=True,
    )
    para(
        doc,
        "How inDoc Answers Questions Hidden Across A Document Library",
        size=22,
        bold=True,
        color=C_NAVY,
        after=6,
        keep_next=True,
    )
    para(
        doc,
        "Similarity finds the paragraph. A graph finds the pattern.",
        size=13,
        italic=True,
        color=C_PRIMARY,
        after=10,
    )

    para(
        doc,
        "A compliance officer opens inDoc on a Friday afternoon and asks a question any counsel would recognize: which clause assigns late-delivery liability in the Acme Master Services Agreement. Hybrid search finds §8.2. The agent reads it. The citation lands on a page a lawyer can open in the same sitting.",
        after=8,
    )
    rich(
        doc,
        [
            (
                "She asks a second question that sounds just as ordinary, and is not. ",
                {},
            ),
            (
                "Across every vendor contract signed since 2021, which liability patterns keep appearing",
                {"italic": True, "color": C_NAVY},
            ),
            (
                " — the same fee cap, the same mutual waiver, the same silent carve-out for consequential damages. The first question has a home. The second is a distribution. inDoc already answers the first. This briefing is about the second, and about the map we would have to build to answer it without walking a regulated library one file at a time.",
                {},
            ),
        ],
        after=8,
    )

    figure(
        doc,
        figs["case"],
        "Figure 1. The Acme MSA Has A Home For §8.2. The Recurring Six-Month Fee Cap Is A Fact About The Portfolio — Six Of Twelve Agreements, Zero Documents That State The Tally.",
    )

    heading(doc, "01", "Two questions that look the same")
    para(
        doc,
        "inDoc customers do not live in generic collections. They live in regulated libraries: vendor contracts, incident postmortems, SOPs, clinical quality files, board packs, and policy trees. The questions they ask fall into two families that look identical in a search box and are opposites in an architecture.",
    )
    para(
        doc,
        "A located question has an answer in a small number of passages. The wording of the question and the wording of the answer overlap: which service owns the payments retry path; what was Q3 revenue; does the 2024 DPA allow EU subprocessors. Keyword search, vector search, and inDoc’s hybrid fuse were built for this family.",
    )
    para(
        doc,
        "A corpus question is a pattern, a count, a recurrence, or a contradiction. No single file contains it. The useful material is spread across tens or hundreds of documents that may never use the same nouns as the question: which failure causes recur across reviews; where indemnity clauses diverge from standard, and which vendors share the deviation; which quality findings keep reopening after close. These questions punish resemblance. The thing required is not nearby text. It is a structure that was never written down as one paragraph.",
    )
    quote(
        doc,
        "A larger prompt window does not convert a survey into a lookup. It only lets the model read more of whatever retrieval already fetched.",
    )

    heading(doc, "02", "What inDoc already answers")
    para(
        doc,
        "inDoc is not a chat wrapper over a single PDF. It is a private document system: a processing pipeline, two search engines, a conversational path, and an autonomous research agent. Every retrieval hop is bound to the requesting user’s scope. That constraint is the product. Healthcare, legal, and finance deployments do not get a second, quieter path around it.",
    )
    figure(
        doc,
        figs["stack"],
        "Figure 2. Shipped system: ingest, scoped hybrid retrieval, budgeted reasoning, and a governance gate. PostgreSQL remains the record. The agent cannot walk the library.",
    )
    para(
        doc,
        "Elasticsearch holds the words. Qdrant holds a 384-dimension cosine vector. Scores are min-max normalized and fused at equal weight. Scope — role, classification, optional selection — is applied before a result leaves the service. Chat, the agent’s search tool, and MCP search share this path. If both engines are down, the service falls back to a scoped database lookup.",
    )
    para(
        doc,
        "The agent is a ReAct loop, not an open reader. It plans, searches, reads or summarizes at most two high-value documents, compares when the goal is a diff, and finishes with citations. The default budget is six steps; the hard cap is twelve. An unbounded walk of a thousand-document library is latency, cost, and a compliance incident waiting for a mis-routed read. Caps are a product decision. The stack is strong at located questions. It is honest about corpus questions: several targeted searches, one or two deep dives, then stop.",
    )

    heading(doc, "03", "Where resemblance ends")
    para(
        doc,
        "Hybrid search still rests on a resemblance assumption. A query becomes tokens for Elasticsearch and a vector for Qdrant. Documents whose text or embedding sit close to that query rise. For located questions, resemblance is a feature. “Payments retry” appears in the architecture decision that assigned ownership. The fuse returns that file. The citation is checkable.",
    )
    para(
        doc,
        "Corpus questions break the assumption. “Which failure causes recur most often” embeds as a sentence about recurrence. The nearest neighbors are files that happen to say recurring, frequent, or again. The true answer is a histogram over two hundred postmortems. That histogram does not live at a point in vector space. It is a property of the collection. Microsoft’s GraphRAG research named this split local versus global. inDoc’s first question is local. The portfolio question is global.",
    )
    figure(
        doc,
        figs["gap"],
        "Figure 3. Similarity returns vocabulary cousins of “recur.” The question wants a distribution over 214 postmortems. Those are different jobs.",
    )
    para(
        doc,
        "Three common escapes do not close the gap. A larger context window still fetches nearest text; if nearest is the wrong set, the extra tokens are extra noise. Asking the agent to read everything is refused — the budget exists so a regulated library cannot be walked file by file. Calling fluent wrong answers “hallucination” often misstates the failure: the model is obedient, and retrieval handed it the wrong neighborhood. The prose is a faithful summary of an irrelevant set.",
    )
    para(
        doc,
        "inDoc’s hybrid fuse reduces one failure mode — exact legal or clinical terms that embeddings smear — but it does not invent a map of how documents relate. Document-level hybrid search is the right default for “find the MSA.” It is the wrong index for “show me the shared deviation across 400 MSAs.”",
    )

    heading(doc, "04", "The map between files")
    para(
        doc,
        "A knowledge graph stores two kinds of objects. Entities are the nouns the library talks about: vendors, clauses, services, teams, incidents, policies, findings, counterparties, control identifiers. Relationships are typed links: a vendor indemnifies a company; an incident follows a deploy; an SOP supersedes another; a finding reopens a closed one. Both carry a short description. Both keep a pointer to the text they came from. That pointer is the same citation contract inDoc already enforces when chat answers and when the agent finishes.",
    )
    figure(
        doc,
        figs["sentence"],
        "Figure 4. One sentence from INC-2024-031 becomes four objects and three typed links. Every object points back to text unit T-1844. That is the audit.",
    )
    para(
        doc,
        "After thousands of sentences contribute nodes and edges, paths appear that no single document contains. An engineer is named in a design decision. A service is named in a runbook. An incident is described in a third file. The path from that engineer to that incident runs through the service. Similarity search never walks that path. It only returns files that look like the query.",
    )
    figure(
        doc,
        figs["path"],
        "Figure 5. Maya Chen → payments retry path → checkout-api → INC-2024-031. Three documents. One walk. No file states the whole chain.",
    )
    quote(
        doc,
        "Throwing away internal structure at ingest is an expensive default. Chunks are not a map.",
    )
    para(
        doc,
        "inDoc already holds the lexical half in pieces: documents, chunks, Elasticsearch text, Qdrant vectors. It does not yet persist a first-class entity graph. That is a precise description of the current product, not a criticism of it.",
    )

    heading(doc, "05", "How the graph would be built, queried, and paid for")
    para(
        doc,
        "Construction is a pipeline, and most of the money is spent in one stage. Documents are sliced into text units. A model extracts entities and relationships. Descriptions that share a title and type are merged — a service named in two hundred files produces two hundred descriptions that must become one node. Optional claims capture time-bounded facts. The entity graph is clustered into a hierarchy. A community report is written at every level so that a survey of what the library collectively says exists before anyone asks. Public cost notes put extraction at roughly three quarters of indexing spend. Prompt quality is domain-specific: a legal extract prompt is not a clinical extract prompt.",
    )
    para(
        doc,
        "Which community level you query is a product control, not an algorithm name. Overview is cheap and abstract. Exhaustive is thorough and expensive. inDoc would expose those words. Community reports inherit the highest classification of their sources, or they are filtered at query time — there is no third option that is safe. Tenants never share an entity node, even if two hospitals both operate a Pharmacy.",
    )
    figure(
        doc,
        figs["operating"],
        "Figure 6. Six-stage construction, three query paths, three cost profiles. Hybrid search stays. The graph is invited, not crowned.",
    )
    para(
        doc,
        "Local graph search expands from matched entities — text units, reports, neighbors, claims — then ranks into one window. Global search leaves the entity graph alone and map-reduces community reports. A blended mode starts from the most relevant reports and follows up with local search. Ordinary hybrid search remains the right tool when the answer looks like the question. The agent already routes. GraphRAG, if we add it, is another tool: not a new product, not a replacement.",
    )
    table2(
        doc,
        ["When The Question…", "inDoc Uses…"],
        [
            [
                "Looks like a passage, or names a file or clause",
                "Hybrid search, then an optional read",
            ],
            ["Names entities and asks how they connect", "Local graph search"],
            [
                "Asks what recurs, differs, or dominates",
                "Global search over community reports",
            ],
            [
                "Is a structured fact already in metadata",
                "PostgreSQL — not a language model",
            ],
            [
                "Needs something outside the library",
                "Refuse, or an explicit tool with a visible boundary",
            ],
        ],
    )
    para(
        doc,
        "The index is derived. A new contract is not upload-and-forget. For a library that changes every day, rebuild is an operations commitment: jobs, progress, and a clear story for what is stale. Lazy construction — a cheap NLP graph, model spend at query time — is the right trade for some tenants. The reason not to make every tenant lazy is that pre-written community reports have value beyond Q&A. People read them. They can be first-class classified documents.",
    )
    para(
        doc,
        "Two evaluation facts matter more than slogans. Hybrid RAG remains the stronger default for local questions — most of day-one use. GraphRAG’s measured gains are comprehensiveness, diversity of sources, and supporting material. Faithfulness, whether each individual claim is true, has scored in the same band as baseline RAG. GraphRAG is a coverage and structure upgrade. It is not a hallucination killer. inDoc will not sell it as one.",
    )

    heading(doc, "06", "What this means for the product")
    figure(
        doc,
        figs["decision"],
        "Figure 7. Keep the system that is already safe. Add a map when the tenant’s real questions are corpus questions. Confirm every printed claim in a source file.",
    )
    para(
        doc,
        "Default stays hybrid search. Graph work is justified when a tenant’s real questions are corpus questions — portfolio-wide contract deviation, recurring quality findings, multi-year incident themes — and that tenant will pay ingest cost and accept rebuild lag. A community report is a summary of a cluster. The last hop still opens the file. Graph retrieval finds the pattern. Hybrid search and a careful read keep the pattern honest.",
    )

    heading(doc, "07", "What to take from this briefing")
    takeaways = [
        "Located questions have a home in a few passages. Corpus questions are properties of the library.",
        "inDoc already investigates under scope, with a budget and a trace. That is the identity. A graph would be the next index.",
        "Extraction without a pointer back to the paragraph cannot survive an audit. Classification and tenant isolation bind derived reports the same way they bind source files.",
        "Coverage is not faithfulness. Better source diversity does not, on published evidence, make each claim truer by itself.",
    ]
    for item in takeaways:
        p = para(doc, "", after=6)
        r = p.add_run("●   ")
        font_run(r, "Calibri", 11, True, C_PRIMARY)
        r = p.add_run(item)
        font_run(r, "Calibri", 11, False, C_TEXT)

    heading(doc, "08", "Method")
    para(
        doc,
        "Original Shared Oxygen / inDoc writing. Public GraphRAG research is used for concepts, not wording: Microsoft Research (local and global queries, community reports, later lazy-index work); project documentation of indexing phases; published graph-structured retrieval results in customer operations, cited for direction rather than reproduced statistics. Claims about inDoc behavior come from the current codebase — hybrid search, agent tools, document scope, and the processing task. Where a graph is future work, it is labeled that way. Figures are original briefing graphics, not third-party art.",
        size=10,
        color=C_MUTED,
        after=10,
    )
    para(
        doc,
        "© 2026 Shared Oxygen, LLC.  Autonomous, but accountable.",
        size=10,
        bold=True,
        color=C_PRIMARY,
        after=0,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


def render() -> dict[str, Path]:
    ASSET.mkdir(parents=True, exist_ok=True)
    return {
        "case": fig_01_case(),
        "stack": fig_02_stack(),
        "gap": fig_03_gap(),
        "sentence": fig_04_sentence(),
        "path": fig_05_path(),
        "operating": fig_06_operating(),
        "decision": fig_07_decision(),
    }


if __name__ == "__main__":
    figs = render()
    path = build(figs)
    print(path)
    print("bytes", path.stat().st_size)
